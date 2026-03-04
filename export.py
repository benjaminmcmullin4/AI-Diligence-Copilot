"""Export utilities: PDF (fpdf2), DOCX (python-docx), and Markdown.

Generates professional IC memos in multiple formats for download.
"""

from __future__ import annotations

import re
from datetime import datetime, timezone

from config import COLORS, FIRM_NAME, FONT


# ═══════════════════════════════════════════════════════════════════════════
# Markdown Export
# ═══════════════════════════════════════════════════════════════════════════

def generate_memo_markdown(
    company_name: str,
    memo_markdown: str,
    business_model: str = "saas",
) -> str:
    """Wrap raw memo markdown with header and footer."""
    date_str = datetime.now(timezone.utc).strftime("%B %d, %Y")
    bm_label = "SaaS" if business_model == "saas" else "Non-SaaS"

    header = (
        f"---\n"
        f"company: {company_name}\n"
        f"type: Investment Committee Memorandum\n"
        f"business_model: {bm_label}\n"
        f"date: {date_str}\n"
        f"prepared_by: {FIRM_NAME} Diligence Copilot\n"
        f"classification: CONFIDENTIAL\n"
        f"---\n\n"
    )

    footer = (
        f"\n\n---\n\n"
        f"*Confidential \u2014 Prepared by {FIRM_NAME} Diligence Copilot on {date_str}*\n"
    )

    return header + memo_markdown + footer


# ═══════════════════════════════════════════════════════════════════════════
# PDF Export (fpdf2)
# ═══════════════════════════════════════════════════════════════════════════

def _strip_markdown(text: str) -> str:
    """Lightweight Markdown-to-plain-text conversion for PDF cells."""
    # Bold
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    # Italic
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    # Inline code
    text = re.sub(r"`(.+?)`", r"\1", text)
    # Links [text](url) -> text
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def generate_memo_pdf(
    company_name: str,
    memo_html_content: str,
    business_model: str = "saas",
    sector: str = "",
    memo_markdown: str = "",
) -> bytes:
    """Generate a professional PDF from memo content using fpdf2.

    Prefers memo_markdown for rendering (converted to plain text for PDF cells).
    Falls back to memo_html_content stripped of tags if markdown is unavailable.

    Returns PDF as bytes.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError(
            "fpdf2 is not installed. Install with: pip install fpdf2"
        )

    date_str = datetime.now(timezone.utc).strftime("%B %Y")
    bm_label = "SaaS" if business_model == "saas" else "Non-SaaS"
    sector_label = sector or ("Software / SaaS" if business_model == "saas" else "Services")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # ── Title / Header ─────────────────────────────────────────────────
    pdf.set_fill_color(10, 10, 10)  # navy
    pdf.rect(0, 0, 210, 35, "F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_xy(15, 8)
    pdf.cell(0, 10, f"{company_name}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_x(15)
    pdf.cell(
        0, 6,
        f"CONFIDENTIAL \u2014 Investment Committee Memorandum | {bm_label} | {sector_label} | {date_str}",
        new_x="LMARGIN", new_y="NEXT",
    )

    pdf.ln(15)

    # ── Body ───────────────────────────────────────────────────────────
    # Use markdown content if available, otherwise strip HTML
    body_text = memo_markdown or memo_html_content
    if not memo_markdown and memo_html_content:
        # Strip HTML tags for plain-text fallback
        body_text = re.sub(r"<[^>]+>", "", memo_html_content)

    pdf.set_text_color(26, 26, 26)  # text color

    for line in body_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue

        # Headings
        if stripped.startswith("## "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(10, 10, 10)
            pdf.cell(0, 8, _strip_markdown(stripped[3:]), new_x="LMARGIN", new_y="NEXT")
            # Teal underline
            pdf.set_draw_color(26, 188, 156)
            pdf.set_line_width(0.6)
            y = pdf.get_y()
            pdf.line(15, y, 195, y)
            pdf.ln(3)
            pdf.set_text_color(26, 26, 26)
            continue
        elif stripped.startswith("### "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(10, 10, 10)
            pdf.cell(0, 7, _strip_markdown(stripped[4:]), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
            pdf.set_text_color(26, 26, 26)
            continue
        elif stripped.startswith("# "):
            pdf.ln(5)
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(10, 10, 10)
            pdf.cell(0, 9, _strip_markdown(stripped[2:]), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
            pdf.set_text_color(26, 26, 26)
            continue

        # Horizontal rules
        if stripped in ("---", "***", "___"):
            pdf.ln(3)
            pdf.set_draw_color(236, 240, 241)
            pdf.set_line_width(0.3)
            y = pdf.get_y()
            pdf.line(15, y, 195, y)
            pdf.ln(3)
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font("Helvetica", "", 10)
            text = _strip_markdown(stripped[2:])
            pdf.set_x(20)
            pdf.cell(5, 5, "\u2022", new_x="END")
            pdf.multi_cell(0, 5, f"  {text}", new_x="LMARGIN", new_y="NEXT")
            continue

        # Numbered items
        num_match = re.match(r"^(\d+)\.\s(.+)", stripped)
        if num_match:
            pdf.set_font("Helvetica", "", 10)
            text = _strip_markdown(num_match.group(2))
            pdf.set_x(20)
            pdf.cell(8, 5, f"{num_match.group(1)}.", new_x="END")
            pdf.multi_cell(0, 5, f" {text}", new_x="LMARGIN", new_y="NEXT")
            continue

        # Regular paragraph
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, _strip_markdown(stripped), new_x="LMARGIN", new_y="NEXT")

    # ── Footer ─────────────────────────────────────────────────────────
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(119, 119, 119)
    pdf.cell(
        0, 5,
        f"Confidential \u2014 Prepared by {FIRM_NAME} Diligence Copilot on {date_str}",
        new_x="LMARGIN", new_y="NEXT", align="C",
    )

    return pdf.output()


# ═══════════════════════════════════════════════════════════════════════════
# DOCX Export (python-docx)
# ═══════════════════════════════════════════════════════════════════════════

def generate_memo_docx(
    company_name: str,
    memo_markdown: str,
    business_model: str = "saas",
    sector: str = "",
) -> bytes:
    """Generate a professional DOCX from memo markdown using python-docx.

    Returns DOCX as bytes.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Install with: pip install python-docx"
        )

    import io

    date_str = datetime.now(timezone.utc).strftime("%B %d, %Y")
    bm_label = "SaaS" if business_model == "saas" else "Non-SaaS"
    sector_label = sector or ("Software / SaaS" if business_model == "saas" else "Services")

    doc = Document()

    # ── Styles ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # ── Title ──────────────────────────────────────────────────────────
    title = doc.add_heading(company_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

    subtitle = doc.add_paragraph(
        f"CONFIDENTIAL \u2014 Investment Committee Memorandum | "
        f"{bm_label} | {sector_label} | {date_str}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in subtitle.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        run.font.italic = True

    doc.add_paragraph("")  # spacer

    # ── Body ───────────────────────────────────────────────────────────
    for line in memo_markdown.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Headings
        if stripped.startswith("### "):
            h = doc.add_heading(_strip_markdown(stripped[4:]), level=3)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
            continue
        elif stripped.startswith("## "):
            h = doc.add_heading(_strip_markdown(stripped[3:]), level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
            continue
        elif stripped.startswith("# "):
            h = doc.add_heading(_strip_markdown(stripped[2:]), level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
            continue

        # Horizontal rules
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 60)
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = _strip_markdown(stripped[2:])
            p = doc.add_paragraph(text, style="List Bullet")
            continue

        # Numbered items
        num_match = re.match(r"^(\d+)\.\s(.+)", stripped)
        if num_match:
            text = _strip_markdown(num_match.group(2))
            p = doc.add_paragraph(text, style="List Number")
            continue

        # Regular paragraph
        doc.add_paragraph(_strip_markdown(stripped))

    # ── Footer ─────────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer_p = doc.add_paragraph(
        f"Confidential \u2014 Prepared by {FIRM_NAME} Diligence Copilot on {date_str}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        run.font.italic = True

    # ── Serialize to bytes ─────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
